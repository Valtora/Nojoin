"""Render Markdown into a python-docx document as native Word elements.

Meeting notes are stored as Markdown, so a DOCX export is only as good as the
Markdown it can understand. The previous exporter classified each line in
isolation, which recognised headings, bullets and bold and passed everything
else through as literal text: a Markdown table arrived in Word as paragraphs of
pipes and hyphens, and italics, inline code and links arrived as their raw
punctuation.

This module parses properly instead, using ``markdown-it-py`` with the rule set
in ``MARKDOWN_RULES``, which the PDF exporter enables on its own parser too, so
the two exports can never disagree about what counts as a table. Tables become
real ``w:tbl`` elements, lists become Word list styles, and inline marks become
run formatting.
"""

from __future__ import annotations

import re
from typing import Any, Callable, Iterable, Sequence

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from markdown_it import MarkdownIt
from markdown_it.token import Token

# Word ships these styles in the default template. Anything deeper than the
# third level reuses the third, which is what Word itself does on paste.
_BULLET_STYLES = ("List Bullet", "List Bullet 2", "List Bullet 3")
_NUMBER_STYLES = ("List Number", "List Number 2", "List Number 3")
_TABLE_STYLE = "Table Grid"
_CODE_FONT = "Courier New"
_MAX_HEADING_LEVEL = 9

_HTML_TAG_PATTERN = re.compile(r"<[^>]+>")
_LINE_BREAK_PATTERN = re.compile(r"<br\s*/?>", re.IGNORECASE)


# CommonMark plus the two GFM rules the editor can actually produce. Kept as a
# named constant because the PDF exporter enables the same set on its own parser
# instance: the two exports must never disagree about what a table is.
MARKDOWN_RULES = ("table", "strikethrough")


def _build_parser() -> MarkdownIt:
    return MarkdownIt("commonmark").enable(list(MARKDOWN_RULES))


class _InlineState:
    """Formatting marks currently open while walking an inline token stream."""

    def __init__(self) -> None:
        self.bold = False
        self.italic = False
        self.strike = False
        self.link: str | None = None


def _add_hyperlink(paragraph: Any, url: str, text: str) -> None:
    """Append a real Word hyperlink.

    python-docx has no API for this, so the relationship and the ``w:hyperlink``
    element are created by hand. Falling back to plain text keeps a malformed
    URL from failing an entire export.
    """
    try:
        relationship_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
    except Exception:  # noqa: BLE001 - a bad URL must not fail the export
        paragraph.add_run(text)
        return

    hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {})
    hyperlink.set(qn("r:id"), relationship_id)

    run_element = paragraph._p.makeelement(qn("w:r"), {})
    properties = paragraph._p.makeelement(qn("w:rPr"), {})
    style = paragraph._p.makeelement(qn("w:rStyle"), {})
    style.set(qn("w:val"), "Hyperlink")
    properties.append(style)
    run_element.append(properties)

    text_element = paragraph._p.makeelement(qn("w:t"), {})
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _add_text(
    paragraph: Any, text: str, state: _InlineState, *, code: bool = False
) -> None:
    if not text:
        return
    if state.link:
        _add_hyperlink(paragraph, state.link, text)
        return

    run = paragraph.add_run(text)
    run.bold = state.bold
    run.italic = state.italic
    run.font.strike = state.strike
    if code:
        run.font.name = _CODE_FONT
        run.font.size = Pt(9)


_MARK_TOGGLES: dict[str, tuple[str, bool]] = {
    "strong_open": ("bold", True),
    "strong_close": ("bold", False),
    "em_open": ("italic", True),
    "em_close": ("italic", False),
    "s_open": ("strike", True),
    "s_close": ("strike", False),
}


def render_inline(paragraph: Any, children: Iterable[Token] | None) -> None:
    """Write an inline token stream into an existing paragraph."""
    state = _InlineState()

    for token in children or []:
        toggle = _MARK_TOGGLES.get(token.type)
        if toggle is not None:
            setattr(state, toggle[0], toggle[1])
        elif token.type == "text":
            _add_text(paragraph, token.content, state)
        elif token.type == "code_inline":
            _add_text(paragraph, token.content, state, code=True)
        elif token.type == "link_open":
            href = token.attrGet("href")
            state.link = str(href) if href else None
        elif token.type == "link_close":
            state.link = None
        elif token.type in ("softbreak", "hardbreak"):
            paragraph.add_run().add_break()
        elif token.type == "html_inline":
            _render_inline_html(paragraph, token.content, state)


def _render_inline_html(paragraph: Any, content: str, state: _InlineState) -> None:
    """Handle the small amount of inline HTML that Markdown tables rely on.

    A line break inside a table cell has no Markdown spelling, so both the
    editor and the notes prompt use ``<br>``. Any other inline HTML is reduced
    to its text so markup never reaches the page.
    """
    if _LINE_BREAK_PATTERN.fullmatch(content.strip()):
        paragraph.add_run().add_break()
        return
    _add_text(paragraph, _HTML_TAG_PATTERN.sub("", content), state)


class _MarkdownDocxRenderer:
    """Walks a flat markdown-it token stream and emits Word elements."""

    def __init__(self, document: Any, *, heading_offset: int = 0) -> None:
        self._document = document
        self._heading_offset = heading_offset
        self._list_stack: list[str] = []
        self._blockquote_depth = 0
        self._handlers: dict[str, Callable[[Sequence[Token], int], int]] = {
            "heading_open": self._render_heading,
            "paragraph_open": self._render_paragraph,
            "blockquote_open": self._open_blockquote,
            "blockquote_close": self._close_blockquote,
            "fence": self._render_code,
            "code_block": self._render_code,
            "hr": self._render_rule,
            "table_open": self._render_table,
            "bullet_list_open": self._push_bullet_list,
            "ordered_list_open": self._push_number_list,
            "bullet_list_close": self._pop_list,
            "ordered_list_close": self._pop_list,
            "html_block": self._render_html_block,
        }

    def render(self, markdown_text: str) -> None:
        tokens = _build_parser().parse(markdown_text or "")
        index = 0
        while index < len(tokens):
            handler = self._handlers.get(tokens[index].type)
            index = handler(tokens, index) if handler else index + 1

    # -- block handlers; each returns the index to continue from ------------

    def _render_heading(self, tokens: Sequence[Token], index: int) -> int:
        level = int(tokens[index].tag[1:])
        heading = self._document.add_heading(
            "", level=min(level + self._heading_offset, _MAX_HEADING_LEVEL)
        )
        return self._render_inline_at(heading, tokens, index + 1)

    def _render_paragraph(self, tokens: Sequence[Token], index: int) -> int:
        paragraph = self._document.add_paragraph(style=self._paragraph_style())
        return self._render_inline_at(paragraph, tokens, index + 1)

    def _paragraph_style(self) -> str | None:
        # A list inside a quote keeps its list style; the bullet carries more
        # meaning than the indent does.
        if self._list_stack:
            return self._current_list_style()
        return "Quote" if self._blockquote_depth else None

    def _open_blockquote(self, tokens: Sequence[Token], index: int) -> int:
        # Depth is tracked rather than recursed into: the paragraphs inside are
        # rendered by the main loop like any others, just with the Quote style.
        self._blockquote_depth += 1
        return index + 1

    def _close_blockquote(self, tokens: Sequence[Token], index: int) -> int:
        self._blockquote_depth = max(0, self._blockquote_depth - 1)
        return index + 1

    def _render_code(self, tokens: Sequence[Token], index: int) -> int:
        paragraph = self._document.add_paragraph()
        run = paragraph.add_run(tokens[index].content.rstrip("\n"))
        run.font.name = _CODE_FONT
        run.font.size = Pt(9)
        return index + 1

    def _render_rule(self, tokens: Sequence[Token], index: int) -> int:
        paragraph = self._document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run("* * *")
        return index + 1

    def _render_html_block(self, tokens: Sequence[Token], index: int) -> int:
        text = _LINE_BREAK_PATTERN.sub("\n", tokens[index].content)
        text = _HTML_TAG_PATTERN.sub("", text).strip()
        if text:
            self._document.add_paragraph(text)
        return index + 1

    def _render_table(self, tokens: Sequence[Token], index: int) -> int:
        rows, next_index = _collect_table_rows(tokens, index)
        if not rows:
            return next_index

        column_count = max(len(row) for _, row in rows)
        table = self._document.add_table(rows=len(rows), cols=column_count)
        table.style = _TABLE_STYLE
        table.autofit = True

        for row_index, (is_header, cells) in enumerate(rows):
            for column_index in range(column_count):
                cell = table.cell(row_index, column_index)
                paragraph = cell.paragraphs[0]
                if column_index < len(cells):
                    render_inline(paragraph, cells[column_index])
                if is_header:
                    _make_header_cell(paragraph)
            if is_header:
                _repeat_header_row(table.rows[row_index])

        # Word runs a table straight into the next block without this.
        self._document.add_paragraph()
        return next_index

    # -- list state ---------------------------------------------------------

    def _push_bullet_list(self, tokens: Sequence[Token], index: int) -> int:
        self._list_stack.append("bullet")
        return index + 1

    def _push_number_list(self, tokens: Sequence[Token], index: int) -> int:
        self._list_stack.append("number")
        return index + 1

    def _pop_list(self, tokens: Sequence[Token], index: int) -> int:
        if self._list_stack:
            self._list_stack.pop()
        return index + 1

    def _current_list_style(self) -> str:
        styles = _BULLET_STYLES if self._list_stack[-1] == "bullet" else _NUMBER_STYLES
        return styles[min(len(self._list_stack) - 1, len(styles) - 1)]

    def _render_inline_at(
        self, paragraph: Any, tokens: Sequence[Token], index: int
    ) -> int:
        if index < len(tokens) and tokens[index].type == "inline":
            render_inline(paragraph, tokens[index].children)
            return index + 1
        return index


def _collect_table_rows(
    tokens: Sequence[Token], index: int
) -> tuple[list[tuple[bool, list[list[Token]]]], int]:
    """Gather a table's cells as inline token lists, plus the index after it."""
    rows: list[tuple[bool, list[list[Token]]]] = []
    in_header = False
    current: list[list[Token]] | None = None
    cursor = index + 1

    while cursor < len(tokens):
        token = tokens[cursor]
        if token.type == "table_close":
            return rows, cursor + 1
        if token.type == "thead_open":
            in_header = True
        elif token.type == "thead_close":
            in_header = False
        elif token.type == "tr_open":
            current = []
        elif token.type == "tr_close" and current is not None:
            rows.append((in_header, current))
            current = None
        elif token.type in ("th_open", "td_open") and current is not None:
            inline = tokens[cursor + 1] if cursor + 1 < len(tokens) else None
            current.append(
                list(inline.children or [])
                if inline and inline.type == "inline"
                else []
            )
        cursor += 1

    return rows, cursor


def _make_header_cell(paragraph: Any) -> None:
    for run in paragraph.runs:
        run.bold = True


def _repeat_header_row(row: Any) -> None:
    """Mark the row as a header so Word repeats it across page breaks."""
    properties = row._tr.get_or_add_trPr()
    header = properties.makeelement(qn("w:tblHeader"), {})
    properties.append(header)


def render_markdown(
    document: Any, markdown_text: str, *, heading_offset: int = 0
) -> None:
    """Append ``markdown_text`` to ``document`` as native Word elements.

    ``heading_offset`` shifts Markdown heading levels so notes can sit beneath a
    section heading the exporter added itself.
    """
    _MarkdownDocxRenderer(document, heading_offset=heading_offset).render(markdown_text)
