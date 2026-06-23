import html
import logging
import os
import re
from datetime import timedelta
from io import BytesIO
from typing import List, Optional

# Python-docx imports for DOCX
from docx import Document as DocxDocument
from fastapi import HTTPException
from markdown_pdf import MarkdownPdf, Section
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm.attributes import flag_modified
from sqlmodel import select

from backend.celery_app import celery_app
from backend.models.pipeline import SpeakerCorrectionScope
from backend.models.recording import (
    LEGACY_RECORDING_REPROCESS_REQUIRED_DETAIL,
    Recording,
    recording_supports_unified_mutations,
)
from backend.models.recording_public import (
    RecordingSpeakerPublicRead,
)
from backend.models.speaker import RecordingSpeaker
from backend.models.transcript import Transcript
from backend.services.recording_identity_service import get_recording_by_public_id
from backend.utils.config_manager import config_manager

logger = logging.getLogger(__name__)


async def _get_owned_recording(
    db: AsyncSession,
    recording_public_id: str,
    user_id: int,
    *,
    options: tuple | None = None,
) -> Recording:
    recording = await get_recording_by_public_id(
        db,
        recording_public_id,
        user_id=user_id,
        options=options,
    )
    if recording is None:
        raise HTTPException(status_code=404, detail="Recording not found")
    return recording


def _dispatch_meeting_edge_refresh(recording_id: int, *, enabled: bool = True) -> None:
    if not enabled:
        return

    try:
        celery_app.send_task(
            "backend.worker.tasks.refresh_meeting_edge_task",
            args=[recording_id],
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning(
            "Failed to dispatch Meeting Edge refresh for recording %s: %s",
            recording_id,
            exc,
        )


async def _get_recording_transcript(
    db: AsyncSession, recording_id: int
) -> Transcript | None:
    statement = select(Transcript).where(Transcript.recording_id == recording_id)
    result = await db.execute(statement)
    return result.scalar_one_or_none()


def _canonical_transcript_writes_enabled() -> bool:
    return bool(config_manager.get("enable_canonical_transcript_writes", True))


def _require_recording_transcript_mutations_supported(recording: Recording) -> None:
    if recording_supports_unified_mutations(recording):
        return
    raise HTTPException(
        status_code=409, detail=LEGACY_RECORDING_REPROCESS_REQUIRED_DETAIL
    )


def _find_segment_index_by_public_id(
    transcript: Transcript, utterance_id: str
) -> int | None:
    for index, segment in enumerate(transcript.segments or []):
        if str(segment.get("id")) == utterance_id:
            return index
    return None


def _get_segment_revision(segment: dict) -> int:
    return int(segment.get("revision") or 1)


# --- Pydantic Models ---


class TranscriptSegmentTextUpdate(BaseModel):
    text: str


class TranscriptSegmentSpeakerUpdate(BaseModel):
    new_speaker_name: str
    global_speaker_id: Optional[int] = None
    diarization_label: Optional[str] = None


class FindReplaceRequest(BaseModel):
    find_text: str
    replace_text: str
    case_sensitive: bool = False
    use_regex: bool = False


class TranscriptSegmentsUpdate(BaseModel):
    segments: List[dict]


class TranscriptUtteranceTextPatch(BaseModel):
    text: str
    expected_revision: Optional[int] = None


class TranscriptUtteranceSpeakerPatch(BaseModel):
    new_speaker_name: str
    global_speaker_id: Optional[int] = None
    diarization_label: Optional[str] = None
    scope: SpeakerCorrectionScope = (
        SpeakerCorrectionScope.SPEAKER_EVERYWHERE_IN_RECORDING
    )
    expected_revision: Optional[int] = None


class TranscriptUtteranceRead(BaseModel):
    id: str
    start: float
    end: float
    start_ms: int
    end_ms: int
    text: str
    speaker: str
    recording_speaker_id: Optional[int] = None
    state: str
    revision: int
    speaker_state: Optional[str] = None
    segment_source: str
    provisional: bool = False
    speaker_manually_edited: bool = False
    text_manually_edited: bool = False
    speaker_confidence: Optional[float] = None
    text_confidence: Optional[float] = None
    speaker_assignment_source: Optional[str] = None
    speaker_assignment_authority: Optional[str] = None
    updated_at: Optional[str] = None
    overlapping_speakers: List[str] = []


class TranscriptUtteranceListRead(BaseModel):
    recording_id: str
    revision: int
    utterances: List[TranscriptUtteranceRead]
    tombstones: List[str] = []
    speakers: List[RecordingSpeakerPublicRead] = []


class NotesUpdate(BaseModel):
    notes: str


class UserNotesUpdate(BaseModel):
    user_notes: str


class MeetingEdgeFocusUpdate(BaseModel):
    meeting_edge_focus: str


class ChatRequest(BaseModel):
    message: str
    tag_ids: Optional[List[int]] = None


# --- Helper Functions ---


def _build_speaker_map(speakers) -> dict:
    """Build a mapping from diarization label to speaker name."""
    speaker_map = {}
    for rs in speakers:
        name = (
            rs.local_name
            or (rs.global_speaker.name if rs.global_speaker else None)
            or rs.name
            or rs.diarization_label
        )
        speaker_map[rs.diarization_label] = name
    return speaker_map


def _format_transcript_text(segments, speaker_map: dict) -> str:
    """Format transcript segments as text."""
    lines = []
    for seg in segments:
        speaker_label = seg.get("speaker", "Unknown")
        speaker_name = speaker_map.get(speaker_label, speaker_label)
        start = seg.get("start", 0)
        minutes = int(start // 60)
        seconds = int(start % 60)
        time_str = f"[{minutes:02d}:{seconds:02d}]"
        text = seg.get("text", "").strip()
        lines.append(f"{time_str} {speaker_name}: {text}")
    return "\n".join(lines)


def _sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe file output."""
    return "".join(
        [c for c in filename if c.isalpha() or c.isdigit() or c in (" ", "-", "_", ".")]
    ).strip()


def _get_recording_speaker_display_name(recording_speaker: RecordingSpeaker) -> str:
    return (
        recording_speaker.local_name
        or (
            recording_speaker.global_speaker.name
            if recording_speaker.global_speaker
            else None
        )
        or recording_speaker.name
        or recording_speaker.diarization_label
    )


import re2


def _apply_find_replace(
    transcript: Transcript,
    find_text: str,
    replace_text: str,
    case_sensitive: bool = False,
    use_regex: bool = False,
) -> int:
    """
    Apply find and replace to both transcript segments and notes.
    Returns the number of segment replacements made.
    """
    if len(find_text) > 1000:
        raise HTTPException(
            status_code=400, detail="Search pattern is too long (max 1000 characters)"
        )

    total_segment_replacements = 0

    # Path 1: Simple string replacement (Case Sensitive + No Regex)
    # this is the fastest and safest path
    if not use_regex and case_sensitive:
        for segment in transcript.segments:
            if find_text in segment["text"]:
                new_text = segment["text"].replace(find_text, replace_text)
                if new_text != segment["text"]:
                    segment["text"] = new_text
                    segment["text_manually_edited"] = True
                    total_segment_replacements += 1

        if transcript.notes and find_text in transcript.notes:
            transcript.notes = transcript.notes.replace(find_text, replace_text)

    # Path 2: Regex-based replacement (Case Insensitive OR Explicit Regex)
    else:
        # Path 2: Regex-based replacement (Case Insensitive OR Explicit Regex)

        # google-re2 does not support flags arg in compile. We must use inline flags.
        # e.g. (?i) for ignore case.

        prefix = ""
        if not case_sensitive:
            prefix = "(?i)"

        if use_regex:
            pattern = prefix + find_text
        else:
            # Escape the text, then prepend flag
            pattern = prefix + re2.escape(find_text)

        try:
            regex = re2.compile(pattern)
        except re2.error:
            # Invalid regex provided by user
            raise HTTPException(status_code=400, detail="Invalid regular expression")

        # Replace in transcript segments
        for segment in transcript.segments:
            # google-re2 python bindings support subn.
            new_text, count = regex.subn(replace_text, segment["text"])
            if count > 0:
                segment["text"] = new_text
                segment["text_manually_edited"] = True
                total_segment_replacements += 1

        # Replace in notes
        if transcript.notes:
            transcript.notes = regex.sub(replace_text, transcript.notes)

    if total_segment_replacements > 0:
        flag_modified(transcript, "segments")
        # Reconstruct full text
        full_text = " ".join([s["text"] for s in transcript.segments])
        transcript.text = full_text

    return total_segment_replacements


def _parse_markdown_line(line: str) -> dict:
    """
    Parse a markdown line to identify type (heading, list, paragraph) and content.
    Basic parser for common notes format, handling indentation.
    """
    stripped_line = line.strip()
    if not stripped_line:
        return {"type": "empty", "content": ""}

    # Counts leading spaces to determine indent level (2 spaces = 1 level).
    leading_spaces = len(line) - len(line.lstrip(" "))
    indent_level = leading_spaces // 2

    if stripped_line.startswith("#"):
        level = len(stripped_line.split(" ")[0])
        content = stripped_line.lstrip("#").strip()
        return {"type": "heading", "level": level, "content": content}

    # Handle indent levels
    if stripped_line.startswith("- ") or stripped_line.startswith("* "):
        content = stripped_line[2:].strip()
        return {"type": "list_item", "content": content, "indent": indent_level}

    # Simple numbered list detection (1. , 2. )
    if re.match(r"^\d+\.\s", stripped_line):
        content = re.sub(r"^\d+\.\s", "", stripped_line).strip()
        return {"type": "list_item", "content": content, "indent": indent_level}

    return {"type": "paragraph", "content": stripped_line}


def _generate_full_markdown(
    recording: Recording,
    transcript: Transcript,
    include_transcript: bool,
    include_notes: bool,
    segments=None,
    effective_duration=None,
) -> str:
    """Generate a single markdown string for the entire export."""
    if segments is None:
        segments = transcript.segments
    if effective_duration is None:
        effective_duration = recording.duration_seconds
    md_lines = []

    # Header
    md_lines.append(f"# {html.escape(recording.name)}")
    md_lines.append("")

    date_str = (
        recording.created_at.strftime("%B %d, %Y at %I:%M %p")
        if recording.created_at
        else "Unknown Date"
    )
    md_lines.append(f"**Date:** {date_str}")
    md_lines.append("")

    duration_str = (
        str(timedelta(seconds=int(effective_duration)))
        if effective_duration
        else "Unknown"
    )
    md_lines.append(f"**Duration:** {duration_str}")
    md_lines.append("")

    if recording.speakers:
        speaker_names = []
        for s in recording.speakers:
            name = (
                s.local_name
                or (s.global_speaker.name if s.global_speaker else None)
                or s.name
                or s.diarization_label
            )
            speaker_names.append(html.escape(name))
        md_lines.append(f"**Speakers:** {', '.join(speaker_names)}")
        md_lines.append("")

    md_lines.append("---")
    md_lines.append("")

    # Notes
    if include_notes and transcript.notes:
        md_lines.append("## Meeting Notes")
        md_lines.append("")
        md_lines.append(transcript.notes)
        md_lines.append("")
        md_lines.append("---")
        md_lines.append("")

    # Transcript
    if include_transcript and segments:
        md_lines.append("## Transcript")
        md_lines.append("")
        speaker_map = _build_speaker_map(recording.speakers)

        for seg in segments:
            speaker_label = seg.get("speaker", "Unknown")
            speaker_name = speaker_map.get(speaker_label, speaker_label)
            text = seg.get("text", "").strip()

            start = seg.get("start", 0)
            minutes = int(start // 60)
            seconds = int(start % 60)
            time_str = f"[{minutes:02d}:{seconds:02d}]"

            # Use bold for speaker to make it visually distinct
            md_lines.append(f"**{time_str} {html.escape(speaker_name)}**")
            md_lines.append(f"{html.escape(text)}")
            md_lines.append("")

    return "\n".join(md_lines)


def _generate_pdf_export(
    recording: Recording,
    transcript: Transcript,
    include_transcript: bool,
    include_notes: bool,
    segments=None,
    effective_duration=None,
) -> bytes:
    """Generate a PDF export using markdown-pdf."""
    markdown_content = _generate_full_markdown(
        recording,
        transcript,
        include_transcript,
        include_notes,
        segments=segments,
        effective_duration=effective_duration,
    )

    pdf = MarkdownPdf(toc_level=2)
    css = "body { font-family: Helvetica, sans-serif; }"
    pdf.add_section(Section(markdown_content), user_css=css)

    # Use a temporary file since the library requires a file path string for saving.

    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        pdf.save(tmp_path)
        with open(tmp_path, "rb") as f:
            pdf_bytes = f.read()
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    return pdf_bytes


def _generate_docx_export(
    recording: Recording,
    transcript: Transcript,
    include_transcript: bool,
    include_notes: bool,
    segments=None,
    effective_duration=None,
) -> bytes:
    """Generate a DOCX export."""
    if segments is None:
        segments = transcript.segments
    if effective_duration is None:
        effective_duration = recording.duration_seconds
    buffer = BytesIO()
    doc = DocxDocument()

    # --- Header ---
    doc.add_heading(recording.name, 0)

    # Metadata
    p = doc.add_paragraph()
    date_str = (
        recording.created_at.strftime("%B %d, %Y at %I:%M %p")
        if recording.created_at
        else "Unknown Date"
    )
    p.add_run(f"Date: {date_str}")

    p = doc.add_paragraph()
    duration_str = (
        str(timedelta(seconds=int(effective_duration)))
        if effective_duration
        else "Unknown"
    )
    p.add_run(f"Duration: {duration_str}")

    if recording.speakers:
        speaker_names = []
        for s in recording.speakers:
            name = (
                s.local_name
                or (s.global_speaker.name if s.global_speaker else None)
                or s.name
                or s.diarization_label
            )
            speaker_names.append(name)
        p = doc.add_paragraph()
        p.add_run(f"Speakers: {', '.join(speaker_names)}")

    doc.add_paragraph()  # Spacer

    # --- Notes ---
    if include_notes and transcript.notes:
        doc.add_heading("Meeting Notes", level=1)

        for line in transcript.notes.split("\n"):
            parsed = _parse_markdown_line(line)
            content = parsed["content"]

            # Basic bold handling for DOCX
            def add_formatted_run(paragraph, text):
                parts = re.split(r"(\*\*.*?\*\*)", text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    else:
                        paragraph.add_run(part)

            if parsed["type"] == "heading":
                doc.add_heading(content, level=min(parsed["level"] + 1, 9))
            elif parsed["type"] == "list_item":
                # Map indent level to docx styles
                # 'List Bullet', 'List Bullet 2', 'List Bullet 3'
                indent = parsed.get("indent", 0)
                if indent == 0:
                    style = "List Bullet"
                else:
                    style = f"List Bullet {min(indent + 1, 3)}"

                try:
                    p = doc.add_paragraph(style=style)
                except KeyError:
                    # Fallback if style doesn't exist
                    p = doc.add_paragraph(style="List Bullet")
                    # Could modify p.paragraph_format.left_indent manually if needed

                add_formatted_run(p, content)
            elif parsed["type"] == "paragraph":
                p = doc.add_paragraph()
                add_formatted_run(p, content)

        if include_transcript:
            doc.add_page_break()

    # --- Transcript ---
    if include_transcript and segments:
        doc.add_heading("Transcript", level=1)

        speaker_map = _build_speaker_map(recording.speakers)

        for seg in segments:
            speaker_label = seg.get("speaker", "Unknown")
            speaker_name = speaker_map.get(speaker_label, speaker_label)
            text = seg.get("text", "").strip()

            start = seg.get("start", 0)
            minutes = int(start // 60)
            seconds = int(start % 60)
            time_str = f"[{minutes:02d}:{seconds:02d}]"

            p = doc.add_paragraph()
            # Speaker and Timestamp bold
            run = p.add_run(f"{time_str} {speaker_name}: ")
            run.bold = True
            p.add_run(text)

    doc.save(buffer)
    return buffer.getvalue()
