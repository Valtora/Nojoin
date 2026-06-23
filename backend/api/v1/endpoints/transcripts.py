import html
import json
import logging
import os
import re
import traceback
import uuid
from datetime import timedelta
from io import BytesIO
from typing import List, Literal, Optional

# Python-docx imports for DOCX
from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException, Query, Response
from fastapi.responses import StreamingResponse

# ReportLab imports for PDF removed
# from reportlab.lib import colors
# ...
# Markdown PDF imports
from markdown_pdf import MarkdownPdf, Section
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload
from sqlalchemy.orm.attributes import flag_modified
from sqlmodel import select
from starlette.concurrency import iterate_in_threadpool

from backend.api.deps import get_current_user, get_db
from backend.api.error_handling import sanitized_http_exception
from backend.celery_app import celery_app
from backend.core.db import async_session_maker
from backend.models.chat import ChatMessage
from backend.models.context_chunk import ContextChunk
from backend.models.pipeline import SpeakerCorrectionScope
from backend.models.recording import (
    LEGACY_RECORDING_REPROCESS_REQUIRED_DETAIL,
    Recording,
    recording_supports_unified_mutations,
)
from backend.models.recording_public import (
    ChatMessagePublicRead,
    RecordingSpeakerPublicRead,
    TranscriptPublicRead,
    serialize_chat_message,
    serialize_recording_speaker,
    serialize_transcript,
)
from backend.models.speaker import GlobalSpeaker, RecordingSpeaker
from backend.models.tag import RecordingTag
from backend.models.transcript import Transcript
from backend.models.user import User
from backend.processing.llm_services import (
    get_llm_backend_with_secondary,
)
from backend.processing.pipeline_metrics import record_pipeline_metric
from backend.services.recording_identity_service import get_recording_by_public_id
from backend.utils.canonical_pipeline import (
    apply_compatibility_segment_replace,
    build_transcript_segments_for_read,
    build_transient_utterance_payloads_from_segments,
    ensure_canonical_backfill,
    filter_recording_speakers_for_public_read,
    list_active_utterances,
    recording_ready_for_canonical_backfill,
    serialize_canonical_delta,
)
from backend.utils.canonical_pipeline import (
    update_utterance_speaker as update_canonical_utterance_speaker,
)
from backend.utils.canonical_pipeline import (
    update_utterance_text as update_canonical_utterance_text,
)
from backend.utils.config_manager import config_manager, is_meeting_edge_enabled
from backend.utils.llm_config import resolve_llm_config_async
from backend.utils.speaker_assignment import (
    matches_speaker_name,
    reconcile_segment_assignment,
    segment_references_label,
)

router = APIRouter()
logger = logging.getLogger(__name__)


async def _get_owned_recording(
    db: AsyncSession,
    recording_public_id: str,
    user_id: int,
    *,
    options: tuple | None = None,
) -> Recording:
    recording = await get_recording_by_public_id(
        db,
        recording_public_id,
        user_id=user_id,
        options=options,
    )
    if recording is None:
        raise HTTPException(status_code=404, detail="Recording not found")
    return recording


def _dispatch_meeting_edge_refresh(recording_id: int, *, enabled: bool = True) -> None:
    if not enabled:
        return

    try:
        celery_app.send_task(
            "backend.worker.tasks.refresh_meeting_edge_task",
            args=[recording_id],
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning(
            "Failed to dispatch Meeting Edge refresh for recording %s: %s",
            recording_id,
            exc,
        )


async def _get_recording_transcript(
    db: AsyncSession, recording_id: int
) -> Transcript | None:
    statement = select(Transcript).where(Transcript.recording_id == recording_id)
    result = await db.execute(statement)
    return result.scalar_one_or_none()


def _canonical_transcript_writes_enabled() -> bool:
    return bool(config_manager.get("enable_canonical_transcript_writes", True))


def _require_recording_transcript_mutations_supported(recording: Recording) -> None:
    if recording_supports_unified_mutations(recording):
        return
    raise HTTPException(
        status_code=409, detail=LEGACY_RECORDING_REPROCESS_REQUIRED_DETAIL
    )


def _find_segment_index_by_public_id(
    transcript: Transcript, utterance_id: str
) -> int | None:
    for index, segment in enumerate(transcript.segments or []):
        if str(segment.get("id")) == utterance_id:
            return index
    return None


def _get_segment_revision(segment: dict) -> int:
    return int(segment.get("revision") or 1)


# --- Pydantic Models ---


class TranscriptSegmentTextUpdate(BaseModel):
    text: str


class TranscriptSegmentSpeakerUpdate(BaseModel):
    new_speaker_name: str
    global_speaker_id: Optional[int] = None
    diarization_label: Optional[str] = None


class FindReplaceRequest(BaseModel):
    find_text: str
    replace_text: str
    case_sensitive: bool = False
    use_regex: bool = False


class TranscriptSegmentsUpdate(BaseModel):
    segments: List[dict]


class TranscriptUtteranceTextPatch(BaseModel):
    text: str
    expected_revision: Optional[int] = None


class TranscriptUtteranceSpeakerPatch(BaseModel):
    new_speaker_name: str
    global_speaker_id: Optional[int] = None
    diarization_label: Optional[str] = None
    scope: SpeakerCorrectionScope = (
        SpeakerCorrectionScope.SPEAKER_EVERYWHERE_IN_RECORDING
    )
    expected_revision: Optional[int] = None


class TranscriptUtteranceRead(BaseModel):
    id: str
    start: float
    end: float
    start_ms: int
    end_ms: int
    text: str
    speaker: str
    recording_speaker_id: Optional[int] = None
    state: str
    revision: int
    speaker_state: Optional[str] = None
    segment_source: str
    provisional: bool = False
    speaker_manually_edited: bool = False
    text_manually_edited: bool = False
    speaker_confidence: Optional[float] = None
    text_confidence: Optional[float] = None
    speaker_assignment_source: Optional[str] = None
    speaker_assignment_authority: Optional[str] = None
    updated_at: Optional[str] = None
    overlapping_speakers: List[str] = []


class TranscriptUtteranceListRead(BaseModel):
    recording_id: str
    revision: int
    utterances: List[TranscriptUtteranceRead]
    tombstones: List[str] = []
    speakers: List[RecordingSpeakerPublicRead] = []


class NotesUpdate(BaseModel):
    notes: str


class UserNotesUpdate(BaseModel):
    user_notes: str


class MeetingEdgeFocusUpdate(BaseModel):
    meeting_edge_focus: str


class ChatRequest(BaseModel):
    message: str
    tag_ids: Optional[List[int]] = None


# --- Helper Functions ---


def _build_speaker_map(speakers) -> dict:
    """Build a mapping from diarization label to speaker name."""
    speaker_map = {}
    for rs in speakers:
        name = (
            rs.local_name
            or (rs.global_speaker.name if rs.global_speaker else None)
            or rs.name
            or rs.diarization_label
        )
        speaker_map[rs.diarization_label] = name
    return speaker_map


def _format_transcript_text(segments, speaker_map: dict) -> str:
    """Format transcript segments as text."""
    lines = []
    for seg in segments:
        speaker_label = seg.get("speaker", "Unknown")
        speaker_name = speaker_map.get(speaker_label, speaker_label)
        start = seg.get("start", 0)
        minutes = int(start // 60)
        seconds = int(start % 60)
        time_str = f"[{minutes:02d}:{seconds:02d}]"
        text = seg.get("text", "").strip()
        lines.append(f"{time_str} {speaker_name}: {text}")
    return "\n".join(lines)


def _sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe file output."""
    return "".join(
        [c for c in filename if c.isalpha() or c.isdigit() or c in (" ", "-", "_", ".")]
    ).strip()


def _get_recording_speaker_display_name(recording_speaker: RecordingSpeaker) -> str:
    return (
        recording_speaker.local_name
        or (
            recording_speaker.global_speaker.name
            if recording_speaker.global_speaker
            else None
        )
        or recording_speaker.name
        or recording_speaker.diarization_label
    )


import re2


def _apply_find_replace(
    transcript: Transcript,
    find_text: str,
    replace_text: str,
    case_sensitive: bool = False,
    use_regex: bool = False,
) -> int:
    """
    Apply find and replace to both transcript segments and notes.
    Returns the number of segment replacements made.
    """
    if len(find_text) > 1000:
        raise HTTPException(
            status_code=400, detail="Search pattern is too long (max 1000 characters)"
        )

    total_segment_replacements = 0

    # Path 1: Simple string replacement (Case Sensitive + No Regex)
    # this is the fastest and safest path
    if not use_regex and case_sensitive:
        for segment in transcript.segments:
            if find_text in segment["text"]:
                new_text = segment["text"].replace(find_text, replace_text)
                if new_text != segment["text"]:
                    segment["text"] = new_text
                    segment["text_manually_edited"] = True
                    total_segment_replacements += 1

        if transcript.notes and find_text in transcript.notes:
            transcript.notes = transcript.notes.replace(find_text, replace_text)

    # Path 2: Regex-based replacement (Case Insensitive OR Explicit Regex)
    else:
        # Path 2: Regex-based replacement (Case Insensitive OR Explicit Regex)

        # google-re2 does not support flags arg in compile. We must use inline flags.
        # e.g. (?i) for ignore case.

        prefix = ""
        if not case_sensitive:
            prefix = "(?i)"

        if use_regex:
            pattern = prefix + find_text
        else:
            # Escape the text, then prepend flag
            pattern = prefix + re2.escape(find_text)

        try:
            regex = re2.compile(pattern)
        except re2.error:
            # Invalid regex provided by user
            raise HTTPException(status_code=400, detail="Invalid regular expression")

        # Replace in transcript segments
        for segment in transcript.segments:
            # google-re2 python bindings support subn.
            new_text, count = regex.subn(replace_text, segment["text"])
            if count > 0:
                segment["text"] = new_text
                segment["text_manually_edited"] = True
                total_segment_replacements += 1

        # Replace in notes
        if transcript.notes:
            transcript.notes = regex.sub(replace_text, transcript.notes)

    if total_segment_replacements > 0:
        flag_modified(transcript, "segments")
        # Reconstruct full text
        full_text = " ".join([s["text"] for s in transcript.segments])
        transcript.text = full_text

    return total_segment_replacements


def _parse_markdown_line(line: str) -> dict:
    """
    Parse a markdown line to identify type (heading, list, paragraph) and content.
    Basic parser for common notes format, handling indentation.
    """
    stripped_line = line.strip()
    if not stripped_line:
        return {"type": "empty", "content": ""}

    # Counts leading spaces to determine indent level (2 spaces = 1 level).
    leading_spaces = len(line) - len(line.lstrip(" "))
    indent_level = leading_spaces // 2

    if stripped_line.startswith("#"):
        level = len(stripped_line.split(" ")[0])
        content = stripped_line.lstrip("#").strip()
        return {"type": "heading", "level": level, "content": content}

    # Handle indent levels
    if stripped_line.startswith("- ") or stripped_line.startswith("* "):
        content = stripped_line[2:].strip()
        return {"type": "list_item", "content": content, "indent": indent_level}

    # Simple numbered list detection (1. , 2. )
    if re.match(r"^\d+\.\s", stripped_line):
        content = re.sub(r"^\d+\.\s", "", stripped_line).strip()
        return {"type": "list_item", "content": content, "indent": indent_level}

    return {"type": "paragraph", "content": stripped_line}


def _generate_full_markdown(
    recording: Recording,
    transcript: Transcript,
    include_transcript: bool,
    include_notes: bool,
    segments=None,
    effective_duration=None,
) -> str:
    """Generate a single markdown string for the entire export."""
    if segments is None:
        segments = transcript.segments
    if effective_duration is None:
        effective_duration = recording.duration_seconds
    md_lines = []

    # Header
    md_lines.append(f"# {html.escape(recording.name)}")
    md_lines.append("")

    date_str = (
        recording.created_at.strftime("%B %d, %Y at %I:%M %p")
        if recording.created_at
        else "Unknown Date"
    )
    md_lines.append(f"**Date:** {date_str}")
    md_lines.append("")

    duration_str = (
        str(timedelta(seconds=int(effective_duration)))
        if effective_duration
        else "Unknown"
    )
    md_lines.append(f"**Duration:** {duration_str}")
    md_lines.append("")

    if recording.speakers:
        speaker_names = []
        for s in recording.speakers:
            name = (
                s.local_name
                or (s.global_speaker.name if s.global_speaker else None)
                or s.name
                or s.diarization_label
            )
            speaker_names.append(html.escape(name))
        md_lines.append(f"**Speakers:** {', '.join(speaker_names)}")
        md_lines.append("")

    md_lines.append("---")
    md_lines.append("")

    # Notes
    if include_notes and transcript.notes:
        md_lines.append("## Meeting Notes")
        md_lines.append("")
        md_lines.append(transcript.notes)
        md_lines.append("")
        md_lines.append("---")
        md_lines.append("")

    # Transcript
    if include_transcript and segments:
        md_lines.append("## Transcript")
        md_lines.append("")
        speaker_map = _build_speaker_map(recording.speakers)

        for seg in segments:
            speaker_label = seg.get("speaker", "Unknown")
            speaker_name = speaker_map.get(speaker_label, speaker_label)
            text = seg.get("text", "").strip()

            start = seg.get("start", 0)
            minutes = int(start // 60)
            seconds = int(start % 60)
            time_str = f"[{minutes:02d}:{seconds:02d}]"

            # Use bold for speaker to make it visually distinct
            md_lines.append(f"**{time_str} {html.escape(speaker_name)}**")
            md_lines.append(f"{html.escape(text)}")
            md_lines.append("")

    return "\n".join(md_lines)


def _generate_pdf_export(
    recording: Recording,
    transcript: Transcript,
    include_transcript: bool,
    include_notes: bool,
    segments=None,
    effective_duration=None,
) -> bytes:
    """Generate a PDF export using markdown-pdf."""
    markdown_content = _generate_full_markdown(
        recording,
        transcript,
        include_transcript,
        include_notes,
        segments=segments,
        effective_duration=effective_duration,
    )

    pdf = MarkdownPdf(toc_level=2)
    css = "body { font-family: Helvetica, sans-serif; }"
    pdf.add_section(Section(markdown_content), user_css=css)

    # Use a temporary file since the library requires a file path string for saving.

    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        pdf.save(tmp_path)
        with open(tmp_path, "rb") as f:
            pdf_bytes = f.read()
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    return pdf_bytes


def _generate_docx_export(
    recording: Recording,
    transcript: Transcript,
    include_transcript: bool,
    include_notes: bool,
    segments=None,
    effective_duration=None,
) -> bytes:
    """Generate a DOCX export."""
    if segments is None:
        segments = transcript.segments
    if effective_duration is None:
        effective_duration = recording.duration_seconds
    buffer = BytesIO()
    doc = DocxDocument()

    # --- Header ---
    doc.add_heading(recording.name, 0)

    # Metadata
    p = doc.add_paragraph()
    date_str = (
        recording.created_at.strftime("%B %d, %Y at %I:%M %p")
        if recording.created_at
        else "Unknown Date"
    )
    p.add_run(f"Date: {date_str}")

    p = doc.add_paragraph()
    duration_str = (
        str(timedelta(seconds=int(effective_duration)))
        if effective_duration
        else "Unknown"
    )
    p.add_run(f"Duration: {duration_str}")

    if recording.speakers:
        speaker_names = []
        for s in recording.speakers:
            name = (
                s.local_name
                or (s.global_speaker.name if s.global_speaker else None)
                or s.name
                or s.diarization_label
            )
            speaker_names.append(name)
        p = doc.add_paragraph()
        p.add_run(f"Speakers: {', '.join(speaker_names)}")

    doc.add_paragraph()  # Spacer

    # --- Notes ---
    if include_notes and transcript.notes:
        doc.add_heading("Meeting Notes", level=1)

        for line in transcript.notes.split("\n"):
            parsed = _parse_markdown_line(line)
            content = parsed["content"]

            # Basic bold handling for DOCX
            def add_formatted_run(paragraph, text):
                parts = re.split(r"(\*\*.*?\*\*)", text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    else:
                        paragraph.add_run(part)

            if parsed["type"] == "heading":
                doc.add_heading(content, level=min(parsed["level"] + 1, 9))
            elif parsed["type"] == "list_item":
                # Map indent level to docx styles
                # 'List Bullet', 'List Bullet 2', 'List Bullet 3'
                indent = parsed.get("indent", 0)
                if indent == 0:
                    style = "List Bullet"
                else:
                    style = f"List Bullet {min(indent + 1, 3)}"

                try:
                    p = doc.add_paragraph(style=style)
                except KeyError:
                    # Fallback if style doesn't exist
                    p = doc.add_paragraph(style="List Bullet")
                    # Could modify p.paragraph_format.left_indent manually if needed

                add_formatted_run(p, content)
            elif parsed["type"] == "paragraph":
                p = doc.add_paragraph()
                add_formatted_run(p, content)

        if include_transcript:
            doc.add_page_break()

    # --- Transcript ---
    if include_transcript and segments:
        doc.add_heading("Transcript", level=1)

        speaker_map = _build_speaker_map(recording.speakers)

        for seg in segments:
            speaker_label = seg.get("speaker", "Unknown")
            speaker_name = speaker_map.get(speaker_label, speaker_label)
            text = seg.get("text", "").strip()

            start = seg.get("start", 0)
            minutes = int(start // 60)
            seconds = int(start % 60)
            time_str = f"[{minutes:02d}:{seconds:02d}]"

            p = doc.add_paragraph()
            # Speaker and Timestamp bold
            run = p.add_run(f"{time_str} {speaker_name}: ")
            run.bold = True
            p.add_run(text)

    doc.save(buffer)
    return buffer.getvalue()


# --- Export Endpoint ---


@router.get("/{recording_id}/export")
async def export_content(
    recording_id: str,
    content_type: Literal["transcript", "notes", "both"] = Query(default="transcript"),
    export_format: Literal["txt", "pdf", "docx"] = Query(default="txt"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Export the transcript and/or notes as a file.
    content_type: 'transcript', 'notes', or 'both'
    export_format: 'txt', 'pdf', 'docx'
    """
    logger.info(
        f"Received export request: recording_id={recording_id}, content_type={content_type}, format={export_format}, user={current_user.id}"
    )

    # 1. Fetch Recording with Speakers (for name resolution)
    recording = await _get_owned_recording(
        db,
        recording_id,
        current_user.id,
        options=(
            selectinload(Recording.speakers).options(
                selectinload(RecordingSpeaker.global_speaker)
            ),
        ),
    )

    # Exclude soft-merged speakers from export output
    if recording.speakers:
        recording.speakers = [s for s in recording.speakers if not s.merged_into_id]

    # 2. Fetch Transcript
    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript:
        raise HTTPException(status_code=404, detail="Transcript not found")

    segments = await db.run_sync(
        lambda sync_session: build_transcript_segments_for_read(
            sync_session, recording.id
        )
    )
    effective_duration = recording.duration_seconds

    # 3. Handle Formats
    include_transcript = content_type in ["transcript", "both"]
    include_notes = content_type in ["notes", "both"]

    if include_notes and not transcript.notes:
        if content_type == "notes":
            raise HTTPException(
                status_code=404, detail="No meeting notes available to export"
            )

    # Generate Filename
    timestamp = (
        recording.created_at.strftime("%Y%m%d") if recording.created_at else "export"
    )
    if content_type == "transcript":
        ftype = "Transcript"
    elif content_type == "notes":
        ftype = "Notes"
    else:
        ftype = "FullExport"

    filename = (
        f"{timestamp}_{_sanitize_filename(recording.name)}_{ftype}.{export_format}"
    )

    # 4. Generate Content
    try:
        if export_format == "pdf":
            content = _generate_pdf_export(
                recording,
                transcript,
                include_transcript,
                include_notes,
                segments=segments,
                effective_duration=effective_duration,
            )
            media_type = "application/pdf"
        elif export_format == "docx":
            content = _generate_docx_export(
                recording,
                transcript,
                include_transcript,
                include_notes,
                segments=segments,
                effective_duration=effective_duration,
            )
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:  # txt
            # Text Generation Logic
            speaker_map = _build_speaker_map(recording.speakers)
            sections = []

            # Add Header even for TXT
            sections.append(recording.name)
            if recording.created_at:
                sections.append(
                    f"Date: {recording.created_at.strftime('%B %d, %Y at %I:%M %p')}"
                )
            if effective_duration:
                sections.append(
                    f"Duration: {str(timedelta(seconds=int(effective_duration)))}"
                )
            if recording.speakers:
                s_names = [
                    s.local_name
                    or (s.global_speaker.name if s.global_speaker else None)
                    or s.name
                    or s.diarization_label
                    for s in recording.speakers
                ]
                sections.append(f"Speakers: {', '.join(s_names)}")
            sections.append("=" * 50)
            sections.append("")

            if include_transcript:
                sections.append("Transcript")
                sections.append("-" * 20)
                sections.append(_format_transcript_text(segments, speaker_map))
                sections.append("")

            if include_notes and transcript.notes:
                sections.append("Meeting Notes")
                sections.append("-" * 20)
                sections.append(transcript.notes)
                sections.append("")

            content = "\n".join(sections)
            media_type = "text/plain"
    except Exception as e:  # noqa: BLE001
        logger.error(f"Export generation failed: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Export generation failed")

    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{recording_id}/utterances", response_model=TranscriptUtteranceListRead)
async def get_transcript_utterances(
    recording_id: str,
    after_revision: Optional[int] = Query(default=None, ge=0),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    recording = await _get_owned_recording(db, recording_id, current_user.id)
    transcript = await _get_recording_transcript(db, recording.id)

    if transcript is None:
        raise HTTPException(status_code=404, detail="Transcript not found")

    if not _canonical_transcript_writes_enabled():
        utterances = build_transient_utterance_payloads_from_segments(transcript)
        if after_revision is not None and after_revision >= 0:
            utterances = []

        speakers_result = await db.execute(
            select(RecordingSpeaker)
            .where(RecordingSpeaker.recording_id == recording.id)
            .options(selectinload(RecordingSpeaker.global_speaker))
        )
        speakers = speakers_result.scalars().all()
        speakers = await db.run_sync(
            lambda sync_session: filter_recording_speakers_for_public_read(
                sync_session,
                recording.id,
                speakers,
            )
        )

        return TranscriptUtteranceListRead(
            recording_id=recording.public_id,
            revision=0,
            utterances=[TranscriptUtteranceRead(**payload) for payload in utterances],
            tombstones=[],
            speakers=[
                serialize_recording_speaker(
                    speaker,
                    recording_public_id=recording.public_id,
                )
                for speaker in speakers
            ],
        )

    revision, utterances, tombstones = await db.run_sync(
        lambda sync_session: (
            ensure_canonical_backfill(sync_session, recording.id),
            serialize_canonical_delta(
                sync_session,
                recording.id,
                after_revision=after_revision,
            ),
        )[1]
    )
    await db.commit()
    await db.refresh(transcript)

    if not utterances and transcript.segments:
        utterances = build_transient_utterance_payloads_from_segments(transcript)
        if after_revision is not None and after_revision >= 0:
            utterances = []
    elif after_revision is not None and revision and after_revision >= revision:
        utterances = []
        tombstones = []

    speakers_result = await db.execute(
        select(RecordingSpeaker)
        .where(RecordingSpeaker.recording_id == recording.id)
        .options(selectinload(RecordingSpeaker.global_speaker))
    )
    speakers = speakers_result.scalars().all()
    speakers = await db.run_sync(
        lambda sync_session: filter_recording_speakers_for_public_read(
            sync_session,
            recording.id,
            speakers,
        )
    )

    return TranscriptUtteranceListRead(
        recording_id=recording.public_id,
        revision=revision,
        utterances=[TranscriptUtteranceRead(**payload) for payload in utterances],
        tombstones=tombstones,
        speakers=[
            serialize_recording_speaker(
                speaker,
                recording_public_id=recording.public_id,
            )
            for speaker in speakers
        ],
    )


@router.patch(
    "/{recording_id}/utterances/{utterance_id}/text",
    response_model=TranscriptPublicRead,
)
async def update_transcript_utterance_text(
    recording_id: str,
    utterance_id: str,
    update: TranscriptUtteranceTextPatch,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    recording = await _get_owned_recording(db, recording_id, current_user.id)
    transcript = await _get_recording_transcript(db, recording.id)

    if transcript is None:
        raise HTTPException(status_code=404, detail="Transcript not found")

    _require_recording_transcript_mutations_supported(recording)

    if not _canonical_transcript_writes_enabled():
        segment_index = _find_segment_index_by_public_id(transcript, utterance_id)
        if segment_index is None:
            raise HTTPException(status_code=404, detail="Utterance not found")
        segment = dict(transcript.segments[segment_index])
        if (
            update.expected_revision is not None
            and _get_segment_revision(segment) != update.expected_revision
        ):
            raise HTTPException(status_code=409, detail="Utterance revision conflict")
        return await update_transcript_segment_text(
            recording_id,
            segment_index,
            TranscriptSegmentTextUpdate(text=update.text),
            db,
            current_user,
        )

    await db.run_sync(
        lambda sync_session: ensure_canonical_backfill(sync_session, recording.id)
    )
    if not await db.run_sync(
        lambda sync_session: bool(list_active_utterances(sync_session, recording.id))
    ):
        raise HTTPException(
            status_code=409,
            detail="Canonical utterances are not available for this recording",
        )

    try:
        await db.run_sync(
            lambda sync_session: update_canonical_utterance_text(
                sync_session,
                recording_id=recording.id,
                utterance_public_id=utterance_id,
                text=update.text,
                actor_user_id=current_user.id,
                expected_revision=update.expected_revision,
            )
        )
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc

    await db.commit()
    await db.refresh(transcript)
    record_pipeline_metric(
        stage="transcript_text_correction_applied",
        recording_id=recording.id,
        payload={
            "utterance_id": utterance_id,
            "text_chars": len(update.text),
        },
        log=logger,
    )
    _dispatch_meeting_edge_refresh(
        recording.id,
        enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
    )
    return serialize_transcript(transcript, recording_public_id=recording.public_id)


@router.patch(
    "/{recording_id}/utterances/{utterance_id}/speaker",
    response_model=TranscriptPublicRead,
)
async def update_transcript_utterance_speaker(
    recording_id: str,
    utterance_id: str,
    update: TranscriptUtteranceSpeakerPatch,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    recording = await _get_owned_recording(db, recording_id, current_user.id)
    transcript = await _get_recording_transcript(db, recording.id)

    if transcript is None:
        raise HTTPException(status_code=404, detail="Transcript not found")
    if not update.new_speaker_name.strip():
        raise HTTPException(status_code=400, detail="Speaker name cannot be empty")

    _require_recording_transcript_mutations_supported(recording)

    if not _canonical_transcript_writes_enabled():
        segment_index = _find_segment_index_by_public_id(transcript, utterance_id)
        if segment_index is None:
            raise HTTPException(status_code=404, detail="Utterance not found")
        segment = dict(transcript.segments[segment_index])
        if (
            update.expected_revision is not None
            and _get_segment_revision(segment) != update.expected_revision
        ):
            raise HTTPException(status_code=409, detail="Utterance revision conflict")
        await update_segment_speaker(
            recording_id,
            segment_index,
            TranscriptSegmentSpeakerUpdate(
                new_speaker_name=update.new_speaker_name,
                global_speaker_id=update.global_speaker_id,
                diarization_label=update.diarization_label,
            ),
            db,
            current_user,
        )
        refreshed_transcript = await _get_recording_transcript(db, recording.id)
        if refreshed_transcript is None:
            raise HTTPException(status_code=404, detail="Transcript not found")
        return serialize_transcript(
            refreshed_transcript, recording_public_id=recording.public_id
        )

    await db.run_sync(
        lambda sync_session: ensure_canonical_backfill(sync_session, recording.id)
    )
    if not await db.run_sync(
        lambda sync_session: bool(list_active_utterances(sync_session, recording.id))
    ):
        raise HTTPException(
            status_code=409,
            detail="Canonical utterances are not available for this recording",
        )

    try:
        await db.run_sync(
            lambda sync_session: update_canonical_utterance_speaker(
                sync_session,
                recording_id=recording.id,
                utterance_public_id=utterance_id,
                new_speaker_name=update.new_speaker_name.strip(),
                global_speaker_id=update.global_speaker_id,
                diarization_label=update.diarization_label,
                scope=update.scope,
                actor_user_id=current_user.id,
                expected_revision=update.expected_revision,
            )
        )
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc

    await db.commit()
    await db.refresh(transcript)
    updated_segment = next(
        (
            segment
            for segment in (transcript.segments or [])
            if str(segment.get("id")) == utterance_id
        ),
        None,
    )
    record_pipeline_metric(
        stage="speaker_correction_applied",
        recording_id=recording.id,
        payload={
            "correction_kind": "utterance_speaker",
            "utterance_id": utterance_id,
            "scope": update.scope.value,
            "new_label": updated_segment.get("speaker") if updated_segment else None,
        },
        log=logger,
    )
    _dispatch_meeting_edge_refresh(
        recording.id,
        enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
    )
    return serialize_transcript(transcript, recording_public_id=recording.public_id)


@router.put("/{recording_id}/segments/{segment_index}")
async def update_segment_speaker(
    recording_id: str,
    segment_index: int,
    update: TranscriptSegmentSpeakerUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Update the speaker for a specific transcript segment.
    Also updates the speaker embedding associations using the audio from this segment.
    """
    # 1. Fetch Recording and Transcript
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    # Fetch transcript with segments
    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript:
        raise HTTPException(status_code=404, detail="Transcript not found")

    _require_recording_transcript_mutations_supported(recording)

    if segment_index < 0 or segment_index >= len(transcript.segments):
        raise HTTPException(status_code=400, detail="Invalid segment index")

    if _canonical_transcript_writes_enabled() and (
        transcript.segments[segment_index].get("id")
        or recording_ready_for_canonical_backfill(recording.status)
    ):
        await db.run_sync(
            lambda sync_session: ensure_canonical_backfill(sync_session, recording.id)
        )
        await db.commit()
        await db.refresh(transcript)
        if segment_index < 0 or segment_index >= len(transcript.segments):
            raise HTTPException(status_code=400, detail="Invalid segment index")

        canonical_segment = dict(transcript.segments[segment_index])
        utterance_id = canonical_segment.get("id")
        if not utterance_id:
            raise HTTPException(
                status_code=409, detail="Canonical utterance identifier is unavailable"
            )

        new_speaker_name = update.new_speaker_name.strip()
        if not new_speaker_name:
            raise HTTPException(status_code=400, detail="Speaker name cannot be empty")

        try:
            await db.run_sync(
                lambda sync_session: update_canonical_utterance_speaker(
                    sync_session,
                    recording_id=recording.id,
                    utterance_public_id=str(utterance_id),
                    new_speaker_name=new_speaker_name,
                    global_speaker_id=update.global_speaker_id,
                    diarization_label=update.diarization_label,
                    scope=SpeakerCorrectionScope.UTTERANCE_ONLY,
                    actor_user_id=current_user.id,
                )
            )
        except LookupError as exc:
            raise HTTPException(status_code=404, detail=str(exc)) from exc
        except RuntimeError as exc:
            raise HTTPException(status_code=409, detail=str(exc)) from exc

        await db.commit()
        await db.refresh(transcript)
        refreshed_segment = dict(transcript.segments[segment_index])
        record_pipeline_metric(
            stage="speaker_correction_applied",
            recording_id=recording.id,
            payload={
                "correction_kind": "segment_speaker",
                "segment_index": segment_index,
                "utterance_id": utterance_id,
                "old_label": canonical_segment.get("speaker"),
                "new_label": refreshed_segment.get("speaker"),
                "duration_s": round(
                    float(refreshed_segment.get("end", 0.0))
                    - float(refreshed_segment.get("start", 0.0)),
                    3,
                ),
            },
            log=logger,
        )
        _dispatch_meeting_edge_refresh(
            recording.id,
            enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
        )

        try:
            target_speaker_id = refreshed_segment.get("recording_speaker_id")
            if target_speaker_id is not None:
                start = refreshed_segment["start"]
                end = refreshed_segment["end"]
                duration = end - start
                if duration > 0.5:
                    celery_app.send_task(
                        "backend.worker.tasks.update_speaker_embedding_task",
                        args=[recording.id, start, end, target_speaker_id],
                    )
        except Exception as e:  # noqa: BLE001
            logger.error(f"Failed to dispatch embedding update task: {e}")

        return {"status": "success", "speaker": refreshed_segment.get("speaker")}

    segment = dict(transcript.segments[segment_index])
    old_label = segment.get("speaker")
    new_speaker_name = update.new_speaker_name.strip()

    if not new_speaker_name:
        raise HTTPException(status_code=400, detail="Speaker name cannot be empty")

    # 2. Resolve Target Speaker
    # Determine the diarization_label to assign to the segment.

    target_label = None
    target_recording_speaker = None

    # Check if speaker exists in this recording (by name or global name)
    # Fetch all recording speakers for name comparison
    stmt = (
        select(RecordingSpeaker)
        .where(RecordingSpeaker.recording_id == recording.id)
        .options(selectinload(RecordingSpeaker.global_speaker))
    )
    result = await db.execute(stmt)
    recording_speakers = result.scalars().all()
    target_speaker_id: Optional[int] = None
    current_recording_speaker = next(
        (
            recording_speaker
            for recording_speaker in recording_speakers
            if recording_speaker.diarization_label == old_label
        ),
        None,
    )

    current_global_speaker_id = (
        current_recording_speaker.global_speaker_id
        if current_recording_speaker
        else None
    )
    current_speaker_name = (
        _get_recording_speaker_display_name(current_recording_speaker)
        if current_recording_speaker
        else old_label
    )

    if update.diarization_label == old_label:
        return {"status": "unchanged", "speaker": old_label}

    if (
        update.global_speaker_id is not None
        and current_global_speaker_id == update.global_speaker_id
    ):
        return {"status": "unchanged", "speaker": old_label}

    if (
        update.diarization_label is None
        and update.global_speaker_id is None
        and matches_speaker_name(current_speaker_name, new_speaker_name)
    ):
        return {"status": "unchanged", "speaker": old_label}

    if update.diarization_label is not None:
        target_recording_speaker = next(
            (
                recording_speaker
                for recording_speaker in recording_speakers
                if recording_speaker.diarization_label == update.diarization_label
            ),
            None,
        )

        if not target_recording_speaker:
            raise HTTPException(
                status_code=404, detail="Speaker not found in recording"
            )

        target_label = target_recording_speaker.diarization_label
        target_speaker_id = target_recording_speaker.id

    if target_label is None and update.global_speaker_id is not None:
        stmt = select(GlobalSpeaker).where(
            GlobalSpeaker.id == update.global_speaker_id,
            GlobalSpeaker.user_id == current_user.id,
        )
        result = await db.execute(stmt)
        global_speaker = result.scalar_one_or_none()

        if not global_speaker:
            raise HTTPException(status_code=404, detail="Global speaker not found")

        target_recording_speaker = next(
            (
                rs
                for rs in recording_speakers
                if rs.global_speaker_id == global_speaker.id
            ),
            None,
        )

        if target_recording_speaker:
            target_label = target_recording_speaker.diarization_label
            target_speaker_id = target_recording_speaker.id
        else:
            target_recording_speaker = next(
                (
                    recording_speaker
                    for recording_speaker in recording_speakers
                    if recording_speaker.global_speaker_id is None
                    and (
                        matches_speaker_name(
                            recording_speaker.local_name, new_speaker_name
                        )
                        or matches_speaker_name(
                            recording_speaker.name, new_speaker_name
                        )
                    )
                ),
                None,
            )

            if target_recording_speaker:
                target_recording_speaker.global_speaker_id = global_speaker.id
                target_recording_speaker.local_name = None
                target_recording_speaker.name = None
                db.add(target_recording_speaker)
                await db.flush()
                target_label = target_recording_speaker.diarization_label
                target_speaker_id = target_recording_speaker.id
            else:
                target_label = f"MANUAL_{uuid.uuid4().hex[:8]}"
                target_recording_speaker = RecordingSpeaker(
                    recording_id=recording.id,
                    diarization_label=target_label,
                    global_speaker_id=global_speaker.id,
                    name=None,
                )
                db.add(target_recording_speaker)
                await db.flush()
                target_speaker_id = target_recording_speaker.id

    # Try to find a local match when no global speaker was explicitly selected.
    if target_label is None:
        for rs in recording_speakers:
            if matches_speaker_name(rs.local_name, new_speaker_name):
                target_label = rs.diarization_label
                target_recording_speaker = rs
                target_speaker_id = rs.id
                break

            if matches_speaker_name(rs.name, new_speaker_name):
                target_label = rs.diarization_label
                target_recording_speaker = rs
                target_speaker_id = rs.id
                break

            if matches_speaker_name(
                rs.global_speaker.name if rs.global_speaker else None,
                new_speaker_name,
            ):
                target_label = rs.diarization_label
                target_recording_speaker = rs
                target_speaker_id = rs.id
                break

    if target_label is None:
        target_label = f"MANUAL_{uuid.uuid4().hex[:8]}"
        target_recording_speaker = RecordingSpeaker(
            recording_id=recording.id,
            diarization_label=target_label,
            local_name=new_speaker_name,
            name=None,
        )
        db.add(target_recording_speaker)
        await db.flush()
        target_speaker_id = target_recording_speaker.id

    # 3. Update Transcript Segment
    updated_segments = [dict(entry) for entry in transcript.segments]
    reconcile_segment_assignment(
        updated_segments, segment_index, old_label, target_label
    )
    updated_segments[segment_index]["speaker_manually_edited"] = True
    transcript.segments = updated_segments
    flag_modified(transcript, "segments")
    db.add(transcript)

    # 5. Cleanup Old Speaker (if unused)
    if old_label and old_label != target_label:
        is_used = any(
            segment_references_label(entry, old_label) for entry in transcript.segments
        )

        if not is_used:
            # Delete the RecordingSpeaker entry
            stmt = select(RecordingSpeaker).where(
                RecordingSpeaker.recording_id == recording.id,
                RecordingSpeaker.diarization_label == old_label,
            )
            result = await db.execute(stmt)
            old_speaker_entry = result.scalar_one_or_none()

            if old_speaker_entry:
                await db.delete(old_speaker_entry)
                # Note: We don't delete the GlobalSpeaker, just the local association

    await db.commit()
    record_pipeline_metric(
        stage="speaker_correction_applied",
        recording_id=recording.id,
        payload={
            "correction_kind": "segment_speaker",
            "segment_index": segment_index,
            "old_label": old_label,
            "new_label": target_label,
            "duration_s": round(
                float(segment.get("end", 0.0)) - float(segment.get("start", 0.0)), 3
            ),
        },
        log=logger,
    )
    _dispatch_meeting_edge_refresh(
        recording.id,
        enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
    )

    # 6. Update Embeddings (Active Learning)
    try:
        if target_speaker_id is not None:
            start = segment["start"]
            end = segment["end"]
            duration = end - start

            if duration > 0.5:
                celery_app.send_task(
                    "backend.worker.tasks.update_speaker_embedding_task",
                    args=[
                        recording.id,
                        start,
                        end,
                        target_speaker_id,
                    ],
                )
    except Exception as e:  # noqa: BLE001
        logger.error(f"Failed to dispatch embedding update task: {e}")

    return {"status": "success", "speaker": target_label}


@router.put(
    "/{recording_id}/segments/{segment_index}/text", response_model=TranscriptPublicRead
)
async def update_transcript_segment_text(
    recording_id: str,
    segment_index: int,
    update: TranscriptSegmentTextUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Update the text content of a specific transcript segment.
    """
    # 0. Check Ownership
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    # 1. Fetch Transcript
    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript:
        raise HTTPException(status_code=404, detail="Transcript not found")

    _require_recording_transcript_mutations_supported(recording)

    if segment_index < 0 or segment_index >= len(transcript.segments):
        raise HTTPException(status_code=400, detail="Invalid segment index")

    if _canonical_transcript_writes_enabled() and (
        transcript.segments[segment_index].get("id")
        or recording_ready_for_canonical_backfill(recording.status)
    ):
        await db.run_sync(
            lambda sync_session: ensure_canonical_backfill(sync_session, recording.id)
        )
        await db.commit()
        await db.refresh(transcript)
        if segment_index < 0 or segment_index >= len(transcript.segments):
            raise HTTPException(status_code=400, detail="Invalid segment index")

        utterance_id = transcript.segments[segment_index].get("id")
        if not utterance_id:
            raise HTTPException(
                status_code=409, detail="Canonical utterance identifier is unavailable"
            )

        try:
            await db.run_sync(
                lambda sync_session: update_canonical_utterance_text(
                    sync_session,
                    recording_id=recording.id,
                    utterance_public_id=str(utterance_id),
                    text=update.text,
                    actor_user_id=current_user.id,
                )
            )
        except LookupError as exc:
            raise HTTPException(status_code=404, detail=str(exc)) from exc
        except RuntimeError as exc:
            raise HTTPException(status_code=409, detail=str(exc)) from exc

        await db.commit()
        await db.refresh(transcript)
        record_pipeline_metric(
            stage="transcript_text_correction_applied",
            recording_id=recording.id,
            payload={
                "segment_index": segment_index,
                "utterance_id": utterance_id,
                "text_chars": len(update.text),
            },
            log=logger,
        )
        _dispatch_meeting_edge_refresh(
            recording.id,
            enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
        )
        return serialize_transcript(transcript, recording_public_id=recording.public_id)

    # 2. Update Segment
    updated_segments = [dict(entry) for entry in transcript.segments]
    updated_segments[segment_index]["text"] = update.text
    updated_segments[segment_index]["text_manually_edited"] = True
    transcript.segments = updated_segments
    flag_modified(transcript, "segments")

    # 3. Reconstruct Full Text
    full_text = " ".join([s["text"] for s in transcript.segments])
    transcript.text = full_text

    db.add(transcript)
    await db.commit()
    await db.refresh(transcript)
    record_pipeline_metric(
        stage="transcript_text_correction_applied",
        recording_id=recording.id,
        payload={
            "segment_index": segment_index,
            "text_chars": len(update.text),
        },
        log=logger,
    )
    _dispatch_meeting_edge_refresh(
        recording.id,
        enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
    )

    return serialize_transcript(transcript, recording_public_id=recording.public_id)


@router.post("/{recording_id}/replace", response_model=TranscriptPublicRead)
async def find_and_replace(
    recording_id: str,
    replace_request: FindReplaceRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Find and replace text across the entire transcript AND meeting notes.
    This ensures consistency between the diarized transcript and generated notes.
    """
    # 0. Check Ownership
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    # 1. Fetch Transcript
    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript:
        raise HTTPException(status_code=404, detail="Transcript not found")

    _require_recording_transcript_mutations_supported(recording)

    find_text = replace_request.find_text
    replace_text = replace_request.replace_text
    case_sensitive = replace_request.case_sensitive
    use_regex = replace_request.use_regex

    if not find_text:
        raise HTTPException(status_code=400, detail="Find text cannot be empty")

    # 2. Apply find/replace to both transcript and notes
    _apply_find_replace(transcript, find_text, replace_text, case_sensitive, use_regex)

    if (
        _canonical_transcript_writes_enabled()
        and recording_ready_for_canonical_backfill(recording.status)
    ):
        await db.run_sync(
            lambda sync_session: ensure_canonical_backfill(sync_session, recording.id)
        )
        await db.run_sync(
            lambda sync_session: apply_compatibility_segment_replace(
                sync_session,
                recording_id=recording.id,
                segments=[dict(segment) for segment in (transcript.segments or [])],
            )
        )

    db.add(transcript)
    await db.commit()
    await db.refresh(transcript)
    _dispatch_meeting_edge_refresh(
        recording.id,
        enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
    )

    return serialize_transcript(transcript, recording_public_id=recording.public_id)


@router.put("/{recording_id}/segments", response_model=TranscriptPublicRead)
async def update_transcript_segments(
    recording_id: str,
    update: TranscriptSegmentsUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Bulk update all segments of a transcript.
    Useful for Undo/Redo operations involving multiple segments.
    """
    # 0. Check Ownership
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    # 1. Fetch Transcript
    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript:
        raise HTTPException(status_code=404, detail="Transcript not found")

    _require_recording_transcript_mutations_supported(recording)

    if _canonical_transcript_writes_enabled() and (
        any(segment.get("id") for segment in update.segments)
        or recording_ready_for_canonical_backfill(recording.status)
    ):
        await db.run_sync(
            lambda sync_session: ensure_canonical_backfill(sync_session, recording.id)
        )
        await db.commit()
        await db.refresh(transcript)

        canonical_segments = [dict(segment) for segment in update.segments]
        if canonical_segments and not any(
            segment.get("id") for segment in canonical_segments
        ):
            if len(canonical_segments) == len(transcript.segments):
                for index, segment in enumerate(canonical_segments):
                    segment["id"] = transcript.segments[index].get("id")

        await db.run_sync(
            lambda sync_session: apply_compatibility_segment_replace(
                sync_session,
                recording_id=recording.id,
                segments=canonical_segments,
            )
        )

        await db.commit()
        await db.refresh(transcript)
        _dispatch_meeting_edge_refresh(
            recording.id,
            enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
        )

        return serialize_transcript(transcript, recording_public_id=recording.public_id)

    # 2. Update Segments
    transcript.segments = update.segments
    flag_modified(transcript, "segments")

    # 3. Reconstruct Full Text
    full_text = " ".join([s.get("text", "") for s in transcript.segments])
    transcript.text = full_text

    db.add(transcript)
    await db.commit()
    await db.refresh(transcript)
    _dispatch_meeting_edge_refresh(
        recording.id,
        enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
    )

    return serialize_transcript(transcript, recording_public_id=recording.public_id)


# --- Notes Endpoints ---


@router.get("/{recording_id}/notes")
async def get_notes(
    recording_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Get the meeting notes for a recording.
    """
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript:
        raise HTTPException(status_code=404, detail="Transcript not found")

    return {"notes": transcript.notes}


@router.get("/{recording_id}/user-notes")
async def get_user_notes(
    recording_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Get the user-authored processing notes for a recording.
    """
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    return {"user_notes": transcript.user_notes if transcript else None}


@router.put("/{recording_id}/user-notes")
async def update_user_notes(
    recording_id: str,
    update: UserNotesUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Update the user-authored processing notes for a recording.
    """
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript:
        transcript = Transcript(recording_id=recording.id)

    transcript.user_notes = update.user_notes.strip() or None
    db.add(transcript)
    await db.commit()
    await db.refresh(transcript)
    _dispatch_meeting_edge_refresh(
        recording.id,
        enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
    )

    return {"user_notes": transcript.user_notes, "status": "success"}


@router.put("/{recording_id}/meeting-edge-focus")
async def update_meeting_edge_focus(
    recording_id: str,
    update: MeetingEdgeFocusUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Update the user-provided Meeting Edge focus prompt for a live meeting.
    """
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript:
        transcript = Transcript(recording_id=recording.id)

    transcript.meeting_edge_focus = update.meeting_edge_focus.strip() or None
    db.add(transcript)
    await db.commit()
    await db.refresh(transcript)
    _dispatch_meeting_edge_refresh(
        recording.id,
        enabled=is_meeting_edge_enabled(getattr(current_user, "settings", None)),
    )

    return {
        "meeting_edge_focus": transcript.meeting_edge_focus,
        "status": "success",
    }


@router.put("/{recording_id}/notes")
async def update_notes(
    recording_id: str,
    update: NotesUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Update the meeting notes for a recording.
    """
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript:
        raise HTTPException(status_code=404, detail="Transcript not found")

    transcript.notes = update.notes
    db.add(transcript)
    await db.commit()
    await db.refresh(transcript)

    return {"notes": transcript.notes, "status": "success"}


@router.post("/{recording_id}/notes/generate")
async def generate_notes(
    recording_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Generate meeting notes using the configured LLM provider.
    """
    # 1. Fetch Recording with Speakers
    recording = await _get_owned_recording(
        db,
        recording_id,
        current_user.id,
        options=(
            selectinload(Recording.speakers).options(
                selectinload(RecordingSpeaker.global_speaker)
            ),
        ),
    )

    # 2. Fetch Transcript
    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript = result.scalar_one_or_none()

    if not transcript or not transcript.segments:
        raise HTTPException(status_code=404, detail="Transcript not found or empty")

    if transcript.notes_status == "generating":
        raise HTTPException(
            status_code=409, detail="Meeting notes are already generating."
        )

    llm_config = await resolve_llm_config_async(db, current_user.settings or {})
    missing_llm_config = llm_config.missing_configuration_message()
    if missing_llm_config:
        transcript.notes_status = "error"
        transcript.error_message = missing_llm_config
        db.add(transcript)
        await db.commit()
        raise HTTPException(
            status_code=400,
            detail=f"{missing_llm_config}. Configure an AI provider and model in Settings.",
        )

    # 4. Call Worker Task
    transcript.notes_status = "generating"
    transcript.error_message = None
    db.add(transcript)
    await db.commit()

    task = celery_app.send_task(
        "backend.worker.tasks.generate_notes_task", args=[recording.id]
    )
    from backend.models.task import register_task_ownership

    await register_task_ownership(db, task.id, current_user.id)

    return {
        "status": "success",
        "notes_status": transcript.notes_status,
        "error_message": transcript.error_message,
        "message": "Note generation started",
    }


# --- Chat Endpoints ---


@router.get("/{recording_id}/chat", response_model=List[ChatMessagePublicRead])
async def get_chat_history(
    recording_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Get the chat history for a recording.
    """
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    # Fetch chat messages
    stmt = (
        select(ChatMessage)
        .where(ChatMessage.recording_id == recording.id)
        .order_by(ChatMessage.created_at)
    )
    result = await db.execute(stmt)
    messages = result.scalars().all()

    return [
        serialize_chat_message(message, recording_public_id=recording.public_id)
        for message in messages
    ]


@router.delete("/{recording_id}/chat")
async def clear_chat_history(
    recording_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Clear the chat history for a recording.
    """
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    # Delete all chat messages for this recording
    stmt = select(ChatMessage).where(ChatMessage.recording_id == recording.id)
    result = await db.execute(stmt)
    messages = result.scalars().all()

    for msg in messages:
        await db.delete(msg)

    await db.commit()
    return {"status": "success"}


@router.post("/{recording_id}/chat")
async def chat_with_meeting(
    recording_id: str,
    request: ChatRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Chat with the meeting transcript using LLM (Streaming).
    """
    # 1. Check Ownership & Fetch Data
    recording = await _get_owned_recording(db, recording_id, current_user.id)

    stmt = select(Transcript).where(Transcript.recording_id == recording.id)
    result = await db.execute(stmt)
    transcript_obj = result.scalar_one_or_none()

    if not transcript_obj:
        raise HTTPException(status_code=404, detail="Transcript not found")

    meeting_notes = transcript_obj.notes or ""

    # 2. Get Chat History
    # Retrieve full history; truncation is deferred to the LLM backend if required.
    stmt = (
        select(ChatMessage)
        .where(ChatMessage.recording_id == recording.id)
        .order_by(ChatMessage.created_at)
    )
    result = await db.execute(stmt)
    db_messages = result.scalars().all()

    # Convert to format expected by LLMBackend
    # Google Gemini: {"role": "user"|"model", "parts": [{"text": ...}]}
    # OpenAI/Anthropic: Adapted by backend logic, generally checks for standard roles.
    # Standardizing on Gemini format for internal consistency before backend adaptation.

    formatted_history = []
    for msg in db_messages:
        role = (
            "user" if msg.role == "user" else "model"
        )  # Gemini uses 'model' instead of 'assistant'
        formatted_history.append({"role": role, "parts": [{"text": msg.content}]})

    user_msg = ChatMessage(
        recording_id=recording.id,
        user_id=current_user.id,
        role="user",
        content=request.message,
    )
    db.add(user_msg)
    await db.commit()

    # --- RAG Context Retrieval ---
    context_text = ""
    relevant_chunks = []

    # --- RAG Context Retrieval ---
    context_text = ""
    relevant_chunks = []

    # Always attempt RAG, at least for the current recording
    try:
        # 1. Get embedding for the user query via Celery
        from fastapi.concurrency import run_in_threadpool

        task = celery_app.send_task(
            "backend.worker.tasks.get_text_embedding_task", args=[request.message]
        )
        embeddings = await run_in_threadpool(task.get, timeout=30)
        query_embedding = embeddings[0]

        # 2. Build Query Condition
        if request.tag_ids:
            # Identify relevant recordings from tags
            subquery = select(RecordingTag.recording_id).where(
                RecordingTag.tag_id.in_(request.tag_ids)
            )
            condition = (ContextChunk.recording_id.in_(subquery)) | (
                ContextChunk.recording_id == recording.id
            )
        else:
            # Only search current recording
            condition = ContextChunk.recording_id == recording.id

        # 3. Vector Search
        stmt = (
            select(ContextChunk)
            .where(condition)
            .order_by(ContextChunk.embedding.cosine_distance(query_embedding))
            .limit(5)
        )

        result = await db.execute(stmt)
        relevant_chunks = result.scalars().all()

        if relevant_chunks:
            context_sections = []
            for chunk in relevant_chunks:
                # Fetch recording with speakers for name resolution
                stmt = (
                    select(Recording)
                    .where(Recording.id == chunk.recording_id)
                    .options(
                        selectinload(Recording.speakers).options(
                            selectinload(RecordingSpeaker.global_speaker)
                        )
                    )
                )
                rec_result = await db.execute(stmt)
                rec = rec_result.scalar_one_or_none()

                rec_name = rec.name if rec else f"Recording {chunk.recording_id}"

                content = chunk.content
                # If it's a transcript chunk with speaker info, resolve speaker names
                if chunk.meta and chunk.meta.get("source") == "transcript" and rec:
                    speaker_map = _build_speaker_map(rec.speakers)
                    # Replace raw labels with names
                    # Replace raw diarization labels (e.g. "SPEAKER_XX:") with resolved names.
                    for label, name in speaker_map.items():
                        if label and name and label != name:
                            content = content.replace(f"{label}:", f"{name}:")

                context_sections.append(f"--- From {rec_name} ---\n{content}")

            context_text = "\n\n".join(context_sections)
            logger.info(f"Retrieved {len(relevant_chunks)} context chunks for chat.")

    except Exception as e:  # noqa: BLE001
        logger.error(f"RAG Retrieval failed: {e}")
        # Continue without context rather than failing

    # Augment the final user message with retrieved RAG context.

    augmented_message = request.message
    if context_text:
        augmented_message = f"Context from related meetings/documents:\n{context_text}\n\nUser Question: {request.message}"

    user_settings = current_user.settings or {}

    llm_config = await resolve_llm_config_async(db, user_settings)

    if not llm_config.api_key and llm_config.provider != "ollama":
        raise HTTPException(
            status_code=400,
            detail=f"No API key configured for {llm_config.provider}. Please configure it in settings.",
        )

    try:
        llm_backend = get_llm_backend_with_secondary(llm_config)
    except ValueError as e:
        raise sanitized_http_exception(
            logger=logger,
            status_code=400,
            client_message="Invalid AI configuration.",
            log_message=f"Rejected chat request for recording {recording_id} due to invalid AI configuration.",
            exc=e,
        )
    except Exception as e:  # noqa: BLE001
        logger.error(f"Failed to initialize LLM backend: {e}")
        raise HTTPException(status_code=500, detail="Failed to initialize AI service")

    # 5. Define Streaming Generator
    async def stream_generator():
        full_response = ""
        try:
            generator = llm_backend.ask_question_streaming(
                user_question=augmented_message,
                meeting_notes=meeting_notes,
                diarized_transcript=None,  # Will be fetched inside using recording_id
                conversation_history=formatted_history,
                recording_id=recording.id,
            )

            # Iterate over the generator response asynchronously using threadpool
            # to prevent blocking the asyncio event loop
            async for chunk in iterate_in_threadpool(generator):
                if isinstance(chunk, dict) and chunk.get("type") == "notes_update":
                    yield f"event: notes_update\ndata: {json.dumps({'status': 'success'})}\n\n"
                else:
                    full_response += str(chunk)
                    # Yield SSE format
                    yield f"data: {json.dumps({'token': str(chunk)})}\n\n"

        except Exception as e:  # noqa: BLE001
            logger.error(f"Streaming error: {e}")
            error_msg = str(e).lower()

            # Map common upstream API failures to friendly messages
            if (
                "503" in error_msg
                or "unavailable" in error_msg
                or "overloaded" in error_msg
            ):
                user_msg = "The AI provider is currently experiencing high demand and is unavailable. Please try again later."
            elif (
                "429" in error_msg or "rate limit" in error_msg or "quota" in error_msg
            ):
                user_msg = "You have exceeded your AI provider's rate limit or quota. Please check your billing or try again later."
            elif "timeout" in error_msg or "deadline" in error_msg:
                user_msg = "The AI provider took too long to respond. Please try again."
            elif (
                "context window was exhausted" in error_msg
                or "done_reason=length" in error_msg
            ):
                user_msg = "The Ollama context window was exhausted before a full answer could be generated. Increase the Ollama context window or choose a larger-context model."
            else:
                user_msg = "An internal error occurred while communicating with the AI service. Please try again."

            yield f"data: {json.dumps({'error': user_msg})}\n\n"
            return

        # 6. Save Assistant Response to DB
        try:
            async with async_session_maker() as session:
                assistant_msg = ChatMessage(
                    recording_id=recording.id,
                    user_id=current_user.id,
                    role="assistant",
                    content=full_response,
                )
                session.add(assistant_msg)
                await session.commit()
        except Exception as e:  # noqa: BLE001
            logger.error(f"Failed to save assistant message: {e}")

        yield "data: [DONE]\n\n"

    return StreamingResponse(stream_generator(), media_type="text/event-stream")
