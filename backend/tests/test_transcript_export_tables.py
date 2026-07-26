"""Export-path coverage for tables in meeting notes (issue #136).

Exercises the DOCX and PDF generators end to end rather than the renderer alone,
so a regression in how notes reach them is caught as well.
"""

from datetime import datetime
from io import BytesIO
from types import SimpleNamespace

import fitz
import pytest
from docx import Document

from backend.api.v1.endpoints.transcripts.helpers import (
    _generate_docx_export,
    _generate_pdf_export,
)

NOTES_WITH_TABLES = """## Key Decisions

| ID | Decision | Rationale |
| --- | --- | --- |
| DEC-001 | Use PostgreSQL as the primary persistence layer | Architecture alignment |

## Action Items / Tasks

| ID | Action | Owner | Contributors | Target date | Jira |
| --- | --- | --- | --- | --- | --- |
| ACT-001 | Prepare the architecture diagram | Peter | Anna, Jon | 2026-08-01 | ENG-1421 |
"""


@pytest.fixture
def recording() -> SimpleNamespace:
    return SimpleNamespace(
        id=1,
        name="Architecture Review",
        created_at=datetime(2026, 7, 26, 10, 0),
        duration_seconds=1800,
        speakers=[
            SimpleNamespace(
                local_name="Peter",
                global_speaker=None,
                name=None,
                diarization_label="SPEAKER_00",
                merged_into_id=None,
            )
        ],
    )


@pytest.fixture
def transcript() -> SimpleNamespace:
    return SimpleNamespace(
        notes=NOTES_WITH_TABLES,
        segments=[
            {"speaker": "SPEAKER_00", "start": 0.0, "end": 4.0, "text": "Let us begin."}
        ],
    )


def test_docx_export_writes_native_tables(recording, transcript) -> None:
    payload = _generate_docx_export(
        recording, transcript, include_transcript=False, include_notes=True
    )
    document = Document(BytesIO(payload))

    assert len(document.tables) == 2
    assert len(document.tables[1].columns) == 6
    assert document.tables[1].rows[1].cells[5].text == "ENG-1421"


def test_docx_export_no_longer_emits_pipe_paragraphs(recording, transcript) -> None:
    payload = _generate_docx_export(
        recording, transcript, include_transcript=False, include_notes=True
    )
    document = Document(BytesIO(payload))

    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "| ID |" not in body_text
    assert "| --- |" not in body_text


def test_docx_export_keeps_the_transcript_alongside_tables(
    recording, transcript
) -> None:
    payload = _generate_docx_export(
        recording, transcript, include_transcript=True, include_notes=True
    )
    document = Document(BytesIO(payload))

    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Let us begin." in body_text
    assert "[00:00] Peter: " in body_text
    assert len(document.tables) == 2


def test_pdf_export_draws_table_borders(recording, transcript) -> None:
    payload = _generate_pdf_export(
        recording, transcript, include_transcript=False, include_notes=True
    )

    with fitz.open(stream=payload, filetype="pdf") as document:
        page = document[0]
        text = page.get_text()
        # Cell borders are vector drawings; without the export CSS the table
        # laid out correctly but drew nothing at all.
        assert len(page.get_drawings()) > 0

    assert "ENG-1421" in text
    # Raw Markdown delimiters must not survive into the rendered page.
    assert "| --- |" not in text


def test_pdf_export_handles_notes_without_tables(recording) -> None:
    plain = SimpleNamespace(notes="## Summary\n\nA short summary.", segments=[])

    payload = _generate_pdf_export(
        recording, plain, include_transcript=False, include_notes=True
    )

    with fitz.open(stream=payload, filetype="pdf") as document:
        assert "A short summary." in document[0].get_text()
