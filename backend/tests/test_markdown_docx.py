"""Tests for the Markdown-to-DOCX renderer used by the notes export.

The exporter these cover replaced a line-by-line parser that emitted a Markdown
table into Word as paragraphs of pipes and hyphens (issue #136).
"""

from docx import Document

from backend.utils.markdown_docx import render_markdown

DECISION_TABLE = """## Key Decisions

| ID | Decision | Rationale |
| --- | --- | --- |
| DEC-001 | Use **PostgreSQL** | Architecture alignment |
| DEC-002 | Keep Memgraph | Graph projection |
"""


def _render(markdown_text: str, *, heading_offset: int = 0) -> Document:
    document = Document()
    render_markdown(document, markdown_text, heading_offset=heading_offset)
    return document


def test_markdown_table_becomes_a_native_word_table() -> None:
    document = _render(DECISION_TABLE)

    assert len(document.tables) == 1
    table = document.tables[0]
    assert (len(table.rows), len(table.columns)) == (3, 3)
    assert [cell.text for cell in table.rows[0].cells] == [
        "ID",
        "Decision",
        "Rationale",
    ]
    assert [cell.text for cell in table.rows[2].cells] == [
        "DEC-002",
        "Keep Memgraph",
        "Graph projection",
    ]


def test_table_pipes_never_reach_the_document_body() -> None:
    # The exact failure reported: rows and delimiters exported as paragraphs.
    document = _render(DECISION_TABLE)

    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "|" not in body_text
    assert "---" not in body_text


def test_table_uses_a_bordered_style_with_a_bold_repeating_header() -> None:
    document = _render(DECISION_TABLE)
    table = document.tables[0]

    assert table.style.name == "Table Grid"
    assert table.rows[0].cells[0].paragraphs[0].runs[0].bold is True
    # tblHeader makes Word repeat the header row across a page break.
    assert "tblHeader" in table.rows[0]._tr.xml
    assert "tblHeader" not in table.rows[1]._tr.xml


def test_inline_formatting_survives_inside_a_cell() -> None:
    document = _render(DECISION_TABLE)
    cell = document.tables[0].rows[1].cells[1]

    assert cell.text == "Use PostgreSQL"
    assert any(run.bold for run in cell.paragraphs[0].runs)


def test_line_break_in_a_cell_is_rendered_as_a_break() -> None:
    # Markdown has no in-cell newline, so both the editor and the notes prompt
    # use <br>. A real newline there would split the row.
    document = _render(
        "| Action | Owner |\n| --- | --- |\n| Prep diagram<br>and circulate | Peter |"
    )

    assert document.tables[0].rows[1].cells[0].text == "Prep diagram\nand circulate"


def test_a_six_column_table_keeps_every_column() -> None:
    document = _render(
        "| ID | Action | Owner | Contributors | Target date | Jira |\n"
        "| --- | --- | --- | --- | --- | --- |\n"
        "| ACT-001 | Prepare the diagram | Peter | Anna, Jon | 2026-08-01 | ENG-1421 |"
    )

    table = document.tables[0]
    assert len(table.columns) == 6
    assert table.rows[1].cells[5].text == "ENG-1421"


def test_ragged_rows_are_padded_rather_than_dropped() -> None:
    document = _render("| A | B | C |\n| --- | --- | --- |\n| only one |")

    table = document.tables[0]
    assert len(table.columns) == 3
    assert [cell.text for cell in table.rows[1].cells] == ["only one", "", ""]


def test_headings_lists_and_marks_render_as_word_elements() -> None:
    document = _render(
        "## Detailed Notes\n\n"
        "### Topic\n\n"
        "- Top level\n"
        "  - Nested\n\n"
        "1. First\n\n"
        "Body with **bold**, *italic* and ~~struck~~ text.\n",
        heading_offset=1,
    )

    styles = [
        (paragraph.style.name, paragraph.text)
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    assert ("Heading 3", "Detailed Notes") in styles
    assert ("Heading 4", "Topic") in styles
    assert ("List Bullet", "Top level") in styles
    assert ("List Bullet 2", "Nested") in styles
    assert ("List Number", "First") in styles

    body = next(p for p in document.paragraphs if p.text.startswith("Body with"))
    assert any(run.bold for run in body.runs)
    assert any(run.italic for run in body.runs)
    assert any(run.font.strike for run in body.runs)
    # Strikethrough must be parsed, not passed through as punctuation.
    assert "~~" not in body.text


def test_links_become_real_hyperlinks() -> None:
    document = _render("See [the plan](https://example.com/plan) for detail.")

    paragraph = document.paragraphs[0]
    assert "w:hyperlink" in paragraph._p.xml
    assert "https://example.com/plan" not in paragraph.text


def test_empty_notes_render_nothing() -> None:
    document = _render("")

    assert document.tables == []
    assert all(not paragraph.text.strip() for paragraph in document.paragraphs)
